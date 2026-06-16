"""Concrete Microsoft Word (.docx) export writer."""

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path

from docx import Document as DocxDocument

from src.contexts.studio.application.ports.export_writer import ExportChapter


class DocxExportWriter:
    """Write prepared chapters as a .docx document."""

    def write(
        self,
        path: Path,
        title: str,
        chapters: Iterable[ExportChapter],
    ) -> None:
        """Write ``chapters`` to ``path`` in DOCX format."""
        output = DocxDocument()
        output.add_heading(title, 0)
        for chapter in chapters:
            output.add_heading(chapter.title, level=1)
            for paragraph in chapter.plain_text.split("\n\n"):
                if paragraph.strip():
                    output.add_paragraph(paragraph.strip())
        output.save(str(path))
