from __future__ import annotations

from src.contexts.studio.application.service_common import (
    Any,
    DocumentDto,
    ExportFormat,
    InvalidOperation,
    Iterable,
    NotFound,
    Path,
    Principal,
    RevisionDto,
    StudioRepository,
    _escape_html,
    _export_payload,
    _owner_scopes,
    _plain_text,
    _stream_sha256,
    new_id,
    utcnow,
)

__all__ = ["ExportService"]


class ExportService:
    """Export artifact generation and download."""

    def __init__(
        self,
        repository: StudioRepository,
        *,
        data_dir: Path,
    ) -> None:
        self._repository = repository
        self._data_dir = data_dir

    def export_project(
        self,
        principal: Principal,
        project_id: str,
        *,
        export_format: ExportFormat,
    ) -> dict[str, Any]:
        if export_format not in {"markdown", "docx", "epub"}:
            raise InvalidOperation(f"Unsupported export format: {export_format}")
        owner_id, guest_session_id = _owner_scopes(principal)

        project = self._repository.get_project(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        snapshot = self._repository.get_latest_export_snapshot(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        current_revisions = {
            document.id: document.current_revision_id
            for document in self._repository.list_documents(project_id)
        }
        snapshot_revisions = (
            self._repository.snapshot_revision_map(snapshot.id)
            if snapshot is not None
            else {}
        )
        if snapshot is None or snapshot_revisions != current_revisions:
            snapshot = self._repository.create_snapshot(
                project_id,
                owner_id=owner_id,
                guest_session_id=guest_session_id,
                reason="export",
                now=utcnow(),
            )
        content = [
            (document, revision)
            for document, revision in self._repository.snapshot_content(snapshot.id)
            if document.kind == "chapter"
        ]
        if not content:
            raise InvalidOperation("Create at least one chapter before exporting.")
        export_id = new_id()
        output_dir = self._data_dir / "exports" / project_id
        output_dir.mkdir(parents=True, exist_ok=True)
        suffix = "md" if export_format == "markdown" else export_format
        output_path = output_dir / f"{export_id}.{suffix}"
        if export_format == "markdown":
            self._write_markdown(output_path, project.title, content)
        elif export_format == "docx":
            self._write_docx(output_path, project.title, content)
        else:
            self._write_epub(output_path, project.title, content)
        checksum = _stream_sha256(output_path)
        relative_path = output_path.relative_to(self._data_dir).as_posix()
        export = self._repository.create_export(
            export_id=export_id,
            project_id=project_id,
            snapshot_id=snapshot.id,
            export_format=export_format,
            relative_path=relative_path,
            size_bytes=output_path.stat().st_size,
            checksum_sha256=checksum,
            now=utcnow(),
        )
        return _export_payload(export)

    def list_exports(
        self,
        principal: Principal,
        project_id: str,
    ) -> list[dict[str, Any]]:
        owner_id, guest_session_id = _owner_scopes(principal)
        exports = self._repository.list_exports(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return [_export_payload(item) for item in exports]

    def export_path(
        self,
        principal: Principal,
        project_id: str,
        export_id: str,
    ) -> Path:
        owner_id, guest_session_id = _owner_scopes(principal)
        item = self._repository.get_export(
            project_id,
            export_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        root = self._data_dir.resolve()
        path = (root / item.relative_path).resolve()
        if root not in {path, *path.parents} or not path.is_file():
            raise NotFound("Export file not found.")
        return path

    @staticmethod
    def _write_markdown(
        path: Path,
        title: str,
        content: Iterable[tuple[DocumentDto, RevisionDto]],
    ) -> None:
        parts = [f"# {title}"]
        parts.extend(revision.content_markdown.strip() for _, revision in content)
        path.write_text("\n\n".join(parts).strip() + "\n", encoding="utf-8")

    @staticmethod
    def _write_docx(
        path: Path,
        title: str,
        content: Iterable[tuple[DocumentDto, RevisionDto]],
    ) -> None:
        from docx import Document as DocxDocument

        output = DocxDocument()
        output.add_heading(title, 0)
        for document, revision in content:
            output.add_heading(document.title, level=1)
            for paragraph in _plain_text(revision.content_markdown).split("\n\n"):
                if paragraph.strip():
                    output.add_paragraph(paragraph.strip())
        output.save(str(path))

    @staticmethod
    def _write_epub(
        path: Path,
        title: str,
        content: Iterable[tuple[DocumentDto, RevisionDto]],
    ) -> None:
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier(new_id())
        book.set_title(title)
        book.set_language("en")
        chapters = []
        for index, (document, revision) in enumerate(content, start=1):
            chapter = epub.EpubHtml(
                title=document.title,
                file_name=f"chapter-{index:03d}.xhtml",
                lang="en",
            )
            paragraphs = "".join(
                f"<p>{_escape_html(paragraph)}</p>"
                for paragraph in _plain_text(revision.content_markdown).split("\n\n")
                if paragraph.strip()
            )
            chapter.content = f"<h1>{_escape_html(document.title)}</h1>{paragraphs}"
            book.add_item(chapter)
            chapters.append(chapter)
        book.toc = tuple(chapters)
        book.spine = ["nav", *chapters]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        epub.write_epub(path, book)
