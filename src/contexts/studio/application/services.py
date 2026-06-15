"""Transactional application services for Novel Studio.

The application layer now depends only on the ``StudioRepository`` port and
small injected callables. All SQLAlchemy/model details are encapsulated by the
infrastructure implementation.
"""

from __future__ import annotations

import hashlib
import logging
import re
import secrets
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any, Iterable, TypeVar, cast

import bcrypt
import yaml

from src.contexts.ai.application.ports.text_generation_port import (
    TextGenerationProviderName,
)
from src.contexts.studio.application.ports import (
    DocumentDto,
    ExportDto,
    JobDto,
    ProjectDto,
    ReviewDto,
    RevisionDto,
    SnapshotDto,
    StudioRepository,
    TextGenerationProviderFactory,
)
from src.contexts.studio.domain.exceptions import (
    InvalidOperation,
    NotFound,
    RevisionConflict,
)
from src.contexts.studio.domain.types import DOCUMENT_KINDS, DocumentKind, ExportFormat
from src.contexts.studio.domain.utils import (
    _token_hash,
    _word_count,
    dump_json,
    load_json,
    new_id,
    utcnow,
)

GUEST_TTL = timedelta(hours=24)
SESSION_COOKIE = "novel_studio_session"
CSRF_COOKIE = "novel_studio_csrf"

# Constant dummy bcrypt hash used to keep login timing constant regardless of
# whether the supplied username exists. A fresh salt is generated at import
# time so no hardcoded password literal lives in source control.
_DUMMY_HASH = bcrypt.hashpw(secrets.token_bytes(32), bcrypt.gensalt())

logger = logging.getLogger(__name__)

T = TypeVar("T")


def iso(value: datetime | None) -> str | None:
    """Serialize a datetime to ISO-8601 UTC with ``Z`` suffix."""
    return value.astimezone(UTC).isoformat().replace("+00:00", "Z") if value else None


def _plain_text(markdown: str) -> str:
    text_value = re.sub(r"```.*?```", "", markdown, flags=re.DOTALL)
    text_value = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", text_value)
    text_value = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text_value)
    text_value = re.sub(r"^[#>*+-]+\s*", "", text_value, flags=re.MULTILINE)
    text_value = re.sub(r"[*_`~]", "", text_value)
    return text_value.strip()


def _as_utc(value: datetime) -> datetime:
    return value.replace(tzinfo=UTC) if value.tzinfo is None else value.astimezone(UTC)


def _escape_html(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


# Known mechanical phrases emitted by providers such as DashScope. These are
# stripped or rewritten before the proposal is returned or persisted so the
# manuscript stays in-world.
_FORBIDDEN_TEMPLATE_PHRASES = (
    "revision anchor",
    "the chapter closes",
    "the next scene",
    "first draft",
    "rewritten chapter",
    "focus character",
    "focus_motivation",
    "relationship_status",
    "outline_hook",
)

_FORBIDDEN_TEMPLATE_PATTERN = "|".join(
    re.escape(phrase) for phrase in _FORBIDDEN_TEMPLATE_PHRASES
)
_MECHANICAL_PREAMBLE_RE = re.compile(
    rf"^\s*(?:here(?:'s| is)|below is|sure[,!:]?|certainly[,!:]?|"
    rf"as requested[,!:]?|draft(?:ed)? chapter)\b.*"
    rf"(?:{_FORBIDDEN_TEMPLATE_PATTERN}).*$",
    re.IGNORECASE,
)
_FORBIDDEN_TEMPLATE_REPLACEMENTS: tuple[tuple[re.Pattern[str], str], ...] = (
    (re.compile(r"revision anchor:\s*", re.IGNORECASE), ""),
    (re.compile(r"\bthe chapter closes\b", re.IGNORECASE), "The scene settles"),
    (re.compile(r"\bthe next scene\b", re.IGNORECASE), "What follows"),
    (re.compile(r"\bfirst draft\b", re.IGNORECASE), "opening passage"),
    (re.compile(r"\brewritten chapter\b", re.IGNORECASE), "reworked passage"),
    (re.compile(r"\bfocus character\b", re.IGNORECASE), "central figure"),
    (re.compile(r"\bfocus_motivation\b", re.IGNORECASE), "central motivation"),
    (re.compile(r"\brelationship_status\b", re.IGNORECASE), "relationship state"),
    (re.compile(r"\boutline_hook\b", re.IGNORECASE), "story hook"),
)


def _sanitize_chapter_markdown(markdown: str) -> str:
    """Remove provider preambles and mechanical labels before manuscript storage."""
    cleaned_lines: list[str] = []
    for line in str(markdown).strip().splitlines():
        if _MECHANICAL_PREAMBLE_RE.search(line):
            continue
        cleaned_lines.append(line)
    cleaned = "\n".join(cleaned_lines)
    for pattern, replacement in _FORBIDDEN_TEMPLATE_REPLACEMENTS:
        cleaned = pattern.sub(replacement, cleaned)
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


@dataclass(frozen=True, slots=True)
class Principal:
    """Identity of the caller for authorization and scoping."""

    session_id: str
    kind: str
    owner_id: str | None
    expires_at: datetime | None


def _owner_scopes(principal: Principal) -> tuple[str | None, str | None]:
    """Return (owner_id, guest_session_id) used to scope repository queries."""
    if principal.kind == "owner" and principal.owner_id:
        return principal.owner_id, None
    return None, principal.session_id


# ---------------------------------------------------------------------------
# Payload helpers (DTO -> API dict)
# ---------------------------------------------------------------------------
def _project_payload(
    project: ProjectDto,
    *,
    include_documents: bool = True,
) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "id": project.id,
        "title": project.title,
        "description": project.description,
        "settings": load_json(project.settings_json),
        "import_hash": project.import_hash,
        "created_at": iso(project.created_at),
        "updated_at": iso(project.updated_at),
    }
    if include_documents:
        documents = project.documents or []
        payload["documents"] = [_document_payload(document) for document in documents]
    return payload


def _document_payload(document: DocumentDto) -> dict[str, Any]:
    revision = document.current_revision
    if revision is None:
        raise InvalidOperation("Document has no current revision.")
    return {
        "id": document.id,
        "project_id": document.project_id,
        "kind": document.kind,
        "title": document.title,
        "position": document.position,
        "current_revision_id": revision.id,
        "content_markdown": revision.content_markdown,
        "metadata": load_json(revision.metadata_json),
        "revision_source": revision.source,
        "word_count": _word_count(revision.content_markdown),
        "created_at": iso(document.created_at),
        "updated_at": iso(document.updated_at),
    }


def _revision_payload(revision: RevisionDto) -> dict[str, Any]:
    return {
        "id": revision.id,
        "document_id": revision.document_id,
        "parent_revision_id": revision.parent_revision_id,
        "revision_number": revision.revision_number,
        "content_markdown": revision.content_markdown,
        "metadata": load_json(revision.metadata_json),
        "source": revision.source,
        "word_count": _word_count(revision.content_markdown),
        "created_at": iso(revision.created_at),
    }


def _snapshot_payload(snapshot: SnapshotDto) -> dict[str, Any]:
    return {
        "id": snapshot.id,
        "project_id": snapshot.project_id,
        "reason": snapshot.reason,
        "created_at": iso(snapshot.created_at),
        "documents": [
            {
                "document_id": item.document_id,
                "revision_id": item.revision_id,
                "position": item.position,
            }
            for item in snapshot.documents
        ],
    }


def _review_payload(review: ReviewDto) -> dict[str, Any]:
    return {
        "id": review.id,
        "project_id": review.project_id,
        "snapshot_id": review.snapshot_id,
        "provider": review.provider,
        "model": review.model,
        "summary": review.summary,
        "created_at": iso(review.created_at),
        "issues": [
            {
                "id": issue.id,
                "document_id": issue.document_id,
                "severity": issue.severity,
                "code": issue.code,
                "message": issue.message,
                "suggestion": issue.suggestion,
                "evidence": load_json(issue.evidence_json),
            }
            for issue in review.issues
        ],
    }


def _job_payload(job: JobDto) -> dict[str, Any]:
    return {
        "id": job.id,
        "project_id": job.project_id,
        "document_id": job.document_id,
        "kind": job.kind,
        "operation": job.operation,
        "status": job.status,
        "provider": job.provider,
        "model": job.model,
        "request": load_json(job.request_json),
        "result": load_json(job.result_json),
        "error": job.error,
        "retry_of_job_id": job.retry_of_job_id,
        "created_at": iso(job.created_at),
        "updated_at": iso(job.updated_at),
        "events": [
            {
                "id": event.id,
                "status": event.status,
                "details": load_json(event.details_json),
                "created_at": iso(event.created_at),
            }
            for event in job.events
        ],
    }


def _export_payload(item: ExportDto) -> dict[str, Any]:
    return {
        "id": item.id,
        "project_id": item.project_id,
        "snapshot_id": item.snapshot_id,
        "format": item.format,
        "size_bytes": item.size_bytes,
        "checksum_sha256": item.checksum_sha256,
        "created_at": iso(item.created_at),
        "download_url": f"/api/projects/{item.project_id}/exports/{item.id}/download",
    }


# ---------------------------------------------------------------------------
# Domain services
# ---------------------------------------------------------------------------
class AuthService:
    """Owner configuration and session lifecycle."""

    def __init__(self, repository: StudioRepository) -> None:
        self._repository = repository

    def owner_exists(self) -> bool:
        return self._repository.owner_exists()

    def owner_principal(self, username: str | None = None) -> Principal:
        """Return an operational principal for local CLI maintenance."""
        if username is not None:
            owner = self._repository.get_owner_by_username(username.strip())
        else:
            owner = self._repository.get_first_owner()
        if owner is None:
            raise InvalidOperation("Configure the local owner before importing data.")
        return Principal(
            session_id=f"cli:{owner.id}",
            kind="owner",
            owner_id=owner.id,
            expires_at=None,
        )

    def setup_owner(self, username: str, password: str) -> dict[str, Any]:
        username = username.strip()
        password_bytes = password.encode("utf-8")
        if not username or len(password) < 10 or len(password_bytes) > 72:
            raise InvalidOperation(
                "Username is required and password must be 10-72 UTF-8 bytes."
            )
        if self._repository.owner_exists():
            raise InvalidOperation("The local owner has already been configured.")
        owner = self._repository.create_owner(
            username=username,
            password_hash=bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode("ascii"),
        )
        return {"id": owner.id, "username": owner.username}

    def create_owner_session(
        self,
        username: str,
        password: str,
    ) -> tuple[str, str, Principal]:
        password_bytes = password.encode("utf-8")
        owner = self._repository.get_owner_by_username(username.strip())
        # Always run bcrypt against a real or dummy hash so the timing of the
        # response does not reveal whether the username exists.
        password_hash = (
            owner.password_hash.encode("ascii")
            if owner is not None
            else _DUMMY_HASH
        )
        password_valid = bcrypt.checkpw(password_bytes, password_hash)
        if owner is None or len(password_bytes) > 72 or not password_valid:
            raise InvalidOperation("Invalid username or password.")
        return self._create_session(kind="owner", owner_id=owner.id)

    def create_guest_session(self) -> tuple[str, str, Principal]:
        return self._create_session(
            kind="guest",
            owner_id=None,
            expires_at=utcnow() + GUEST_TTL,
        )

    def _create_session(
        self,
        *,
        kind: str,
        owner_id: str | None,
        expires_at: datetime | None = None,
    ) -> tuple[str, str, Principal]:
        token = secrets.token_urlsafe(36)
        csrf_token = secrets.token_urlsafe(32)
        now = utcnow()
        record = self._repository.create_session(
            kind=kind,
            owner_id=owner_id,
            token_hash=_token_hash(token),
            csrf_token=csrf_token,
            expires_at=expires_at,
            created_at=now,
            last_seen_at=now,
        )
        return token, csrf_token, Principal(record.id, kind, owner_id, expires_at)

    def csrf_token_for_session(self, token_hash: str) -> str | None:
        """Return the CSRF token associated with a session token hash."""
        record = self._repository.get_session_by_token_hash(token_hash)
        return record.csrf_token if record is not None else None

    def principal_from_token(self, token: str | None) -> Principal | None:
        if not token:
            return None
        record = self._repository.get_session_by_token_hash(_token_hash(token))
        if record is None:
            return None
        now = utcnow()
        expires_at = record.expires_at
        if expires_at is not None and _as_utc(expires_at) <= now:
            self._repository.delete_session(record.id)
            return None
        self._repository.update_session_last_seen(record.id, now)
        return Principal(record.id, record.kind, record.owner_id, expires_at)

    def logout(self, token: str | None) -> None:
        if not token:
            return
        record = self._repository.get_session_by_token_hash(_token_hash(token))
        if record is not None:
            self._repository.delete_session(record.id)

    def cleanup_expired_guests(self) -> int:
        return self._repository.delete_expired_guest_sessions(utcnow())


class ProjectService:
    """Project creation and mutation."""

    def __init__(self, repository: StudioRepository) -> None:
        self._repository = repository

    def create_project(
        self,
        principal: Principal,
        *,
        title: str,
        description: str = "",
        create_seed: bool = True,
    ) -> dict[str, Any]:
        title = title.strip()
        if not title:
            raise InvalidOperation("Project title is required.")
        owner_id, guest_session_id = _owner_scopes(principal)
        project = self._repository.create_project(
            owner_id=owner_id,
            guest_session_id=guest_session_id,
            title=title,
            description=description.strip(),
            settings_json=dump_json({"provider": "mock"}),
            now=utcnow(),
            create_seed=create_seed,
        )
        return _project_payload(project)

    def list_projects(self, principal: Principal) -> list[dict[str, Any]]:
        owner_id, guest_session_id = _owner_scopes(principal)
        projects = self._repository.list_projects(
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return [
            _project_payload(project, include_documents=False)
            for project in projects
        ]

    def get_project(self, principal: Principal, project_id: str) -> dict[str, Any]:
        owner_id, guest_session_id = _owner_scopes(principal)
        project = self._repository.get_project(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return _project_payload(project)

    def update_project(
        self,
        principal: Principal,
        project_id: str,
        *,
        title: str | None = None,
        description: str | None = None,
        settings: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        owner_id, guest_session_id = _owner_scopes(principal)
        resolved_title = title.strip() if title is not None and title.strip() else None
        resolved_description = description.strip() if description is not None else None
        settings_json = dump_json(settings) if settings is not None else None
        project = self._repository.update_project(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
            title=resolved_title,
            description=resolved_description,
            settings_json=settings_json,
            now=utcnow(),
        )
        return _project_payload(project)


class DocumentService:
    """Document lifecycle within a project."""

    def __init__(self, repository: StudioRepository) -> None:
        self._repository = repository

    def create_document(
        self,
        principal: Principal,
        project_id: str,
        *,
        kind: DocumentKind,
        title: str,
        content_markdown: str = "",
        position: int | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        if kind not in DOCUMENT_KINDS:
            raise InvalidOperation(f"Unsupported document kind: {kind}")
        owner_id, guest_session_id = _owner_scopes(principal)
        resolved_position = position
        if resolved_position is None:
            resolved_position = self._repository.next_document_position(project_id, kind)
        document = self._repository.create_document(
            project_id=project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
            kind=kind,
            title=title,
            content_markdown=content_markdown,
            position=resolved_position,
            metadata_json=dump_json(metadata or {}),
            source="author",
            now=utcnow(),
        )
        return _document_payload(document)

    def get_document(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
    ) -> dict[str, Any]:
        owner_id, guest_session_id = _owner_scopes(principal)
        document = self._repository.get_document(
            project_id,
            document_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return _document_payload(document)

    def save_document(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
        *,
        content_markdown: str,
        base_revision_id: str | None,
        title: str | None = None,
        metadata: dict[str, Any] | None = None,
        source: str = "author",
    ) -> dict[str, Any]:
        owner_id, guest_session_id = _owner_scopes(principal)
        try:
            document = self._repository.save_document(
                project_id,
                document_id,
                owner_id=owner_id,
                guest_session_id=guest_session_id,
                content_markdown=content_markdown,
                base_revision_id=base_revision_id,
                title=title.strip() if title is not None and title.strip() else None,
                metadata_json=dump_json(metadata or {}),
                source=source,
                now=utcnow(),
            )
        except InvalidOperation as exc:
            if "changed since" in str(exc):
                current_document = self._repository.get_document(
                    project_id,
                    document_id,
                    owner_id=owner_id,
                    guest_session_id=guest_session_id,
                )
                raise RevisionConflict(current_document.current_revision_id) from exc
            raise
        return _document_payload(document)

    def reorder_documents(
        self,
        principal: Principal,
        project_id: str,
        document_ids: list[str],
    ) -> list[dict[str, Any]]:
        owner_id, guest_session_id = _owner_scopes(principal)
        documents = self._repository.reorder_documents(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
            document_ids=document_ids,
            now=utcnow(),
        )
        return [_document_payload(document) for document in documents]

    def search(
        self,
        principal: Principal,
        project_id: str,
        query: str,
    ) -> list[dict[str, Any]]:
        query = query.strip()
        if not query:
            return []
        match_query = f'"{query.replace(chr(34), chr(34) * 2)}"'
        owner_id, guest_session_id = _owner_scopes(principal)
        return self._repository.search_documents(
            project_id,
            match_query,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )


class RevisionService:
    """Revision history and restore operations."""

    def __init__(
        self,
        repository: StudioRepository,
        document_service: DocumentService,
    ) -> None:
        self._repository = repository
        self._document_service = document_service

    def list_revisions(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
    ) -> list[dict[str, Any]]:
        owner_id, guest_session_id = _owner_scopes(principal)
        revisions = self._repository.list_revisions(
            project_id,
            document_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return [_revision_payload(revision) for revision in revisions]

    def restore_revision(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
        revision_id: str,
        *,
        base_revision_id: str | None,
    ) -> dict[str, Any]:
        owner_id, guest_session_id = _owner_scopes(principal)
        revision = self._repository.get_revision(
            project_id,
            document_id,
            revision_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        metadata = cast(dict[str, Any], load_json(revision.metadata_json))
        return self._document_service.save_document(
            principal,
            project_id,
            document_id,
            content_markdown=revision.content_markdown,
            base_revision_id=base_revision_id,
            metadata={**metadata, "restored_from": revision_id},
            source="restore",
        )


class SnapshotService:
    """Project snapshot creation and inspection."""

    def __init__(self, repository: StudioRepository) -> None:
        self._repository = repository

    def create_snapshot(
        self,
        principal: Principal,
        project_id: str,
        *,
        reason: str,
    ) -> dict[str, Any]:
        owner_id, guest_session_id = _owner_scopes(principal)
        snapshot = self._repository.create_snapshot(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
            reason=reason,
            now=utcnow(),
        )
        return _snapshot_payload(snapshot)

    def list_snapshots(
        self,
        principal: Principal,
        project_id: str,
    ) -> list[dict[str, Any]]:
        owner_id, guest_session_id = _owner_scopes(principal)
        snapshots = self._repository.list_snapshots(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return [_snapshot_payload(snapshot) for snapshot in snapshots]


class ReviewService:
    """Editorial review runs."""

    def __init__(self, repository: StudioRepository) -> None:
        self._repository = repository

    def review_project(
        self,
        principal: Principal,
        project_id: str,
        *,
        provider: str = "deterministic",
        model: str = "studio-review-v1",
    ) -> dict[str, Any]:
        owner_id, guest_session_id = _owner_scopes(principal)
        review = self._repository.create_review(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
            provider=provider,
            model=model,
            now=utcnow(),
        )
        return _review_payload(review)

    def list_reviews(
        self,
        principal: Principal,
        project_id: str,
    ) -> list[dict[str, Any]]:
        owner_id, guest_session_id = _owner_scopes(principal)
        reviews = self._repository.list_reviews(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return [_review_payload(review) for review in reviews]


class ExportService:
    """Export artifact generation and download."""

    def __init__(
        self,
        repository: StudioRepository,
        *,
        data_dir: Path,
    ) -> None:
        self._repository = repository
        self._data_dir = data_dir

    def export_project(
        self,
        principal: Principal,
        project_id: str,
        *,
        export_format: ExportFormat,
    ) -> dict[str, Any]:
        if export_format not in {"markdown", "docx", "epub"}:
            raise InvalidOperation(f"Unsupported export format: {export_format}")
        owner_id, guest_session_id = _owner_scopes(principal)

        project = self._repository.get_project(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        snapshot = self._repository.get_latest_export_snapshot(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        current_revisions = {
            document.id: document.current_revision_id
            for document in self._repository.list_documents(project_id)
        }
        snapshot_revisions = (
            self._repository.snapshot_revision_map(snapshot.id)
            if snapshot is not None
            else {}
        )
        if snapshot is None or snapshot_revisions != current_revisions:
            snapshot = self._repository.create_snapshot(
                project_id,
                owner_id=owner_id,
                guest_session_id=guest_session_id,
                reason="export",
                now=utcnow(),
            )
        content = [
            (document, revision)
            for document, revision in self._repository.snapshot_content(snapshot.id)
            if document.kind == "chapter"
        ]
        if not content:
            raise InvalidOperation("Create at least one chapter before exporting.")
        export_id = new_id()
        output_dir = self._data_dir / "exports" / project_id
        output_dir.mkdir(parents=True, exist_ok=True)
        suffix = "md" if export_format == "markdown" else export_format
        output_path = output_dir / f"{export_id}.{suffix}"
        if export_format == "markdown":
            self._write_markdown(output_path, project.title, content)
        elif export_format == "docx":
            self._write_docx(output_path, project.title, content)
        else:
            self._write_epub(output_path, project.title, content)
        checksum = hashlib.sha256(output_path.read_bytes()).hexdigest()
        relative_path = output_path.relative_to(self._data_dir).as_posix()
        export = self._repository.create_export(
            export_id=export_id,
            project_id=project_id,
            snapshot_id=snapshot.id,
            export_format=export_format,
            relative_path=relative_path,
            size_bytes=output_path.stat().st_size,
            checksum_sha256=checksum,
            now=utcnow(),
        )
        return _export_payload(export)

    def list_exports(
        self,
        principal: Principal,
        project_id: str,
    ) -> list[dict[str, Any]]:
        owner_id, guest_session_id = _owner_scopes(principal)
        exports = self._repository.list_exports(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return [_export_payload(item) for item in exports]

    def export_path(
        self,
        principal: Principal,
        project_id: str,
        export_id: str,
    ) -> Path:
        owner_id, guest_session_id = _owner_scopes(principal)
        item = self._repository.get_export(
            project_id,
            export_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        root = self._data_dir.resolve()
        path = (root / item.relative_path).resolve()
        if root not in {path, *path.parents} or not path.is_file():
            raise NotFound("Export file not found.")
        return path

    @staticmethod
    def _write_markdown(
        path: Path,
        title: str,
        content: Iterable[tuple[DocumentDto, RevisionDto]],
    ) -> None:
        parts = [f"# {title}"]
        parts.extend(revision.content_markdown.strip() for _, revision in content)
        path.write_text("\n\n".join(parts).strip() + "\n", encoding="utf-8")

    @staticmethod
    def _write_docx(
        path: Path,
        title: str,
        content: Iterable[tuple[DocumentDto, RevisionDto]],
    ) -> None:
        from docx import Document as DocxDocument

        output = DocxDocument()
        output.add_heading(title, 0)
        for document, revision in content:
            output.add_heading(document.title, level=1)
            for paragraph in _plain_text(revision.content_markdown).split("\n\n"):
                if paragraph.strip():
                    output.add_paragraph(paragraph.strip())
        output.save(str(path))

    @staticmethod
    def _write_epub(
        path: Path,
        title: str,
        content: Iterable[tuple[DocumentDto, RevisionDto]],
    ) -> None:
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier(new_id())
        book.set_title(title)
        book.set_language("en")
        chapters = []
        for index, (document, revision) in enumerate(content, start=1):
            chapter = epub.EpubHtml(
                title=document.title,
                file_name=f"chapter-{index:03d}.xhtml",
                lang="en",
            )
            paragraphs = "".join(
                f"<p>{_escape_html(paragraph)}</p>"
                for paragraph in _plain_text(revision.content_markdown).split("\n\n")
                if paragraph.strip()
            )
            chapter.content = f"<h1>{_escape_html(document.title)}</h1>{paragraphs}"
            book.add_item(chapter)
            chapters.append(chapter)
        book.toc = tuple(chapters)
        book.spine = ["nav", *chapters]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        epub.write_epub(path, book)


class AIService:
    """AI-driven manuscript proposals."""

    def __init__(
        self,
        repository: StudioRepository,
        ai_provider_factory: TextGenerationProviderFactory,
    ) -> None:
        self._repository = repository
        self._ai_provider_factory = ai_provider_factory

    def _load_revision(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
    ) -> tuple[DocumentDto, RevisionDto]:
        owner_id, guest_session_id = _owner_scopes(principal)
        document = self._repository.get_document(
            project_id,
            document_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        revision = document.current_revision
        if revision is None:
            raise InvalidOperation("Document has no current revision.")
        return document, revision

    async def _generate_proposal_text(
        self,
        revision: RevisionDto,
        *,
        operation: str,
        instruction: str,
        provider: str,
        model: str,
    ) -> str:
        from src.contexts.ai.application.ports.text_generation_port import (
            TextGenerationTask,
        )

        task = TextGenerationTask(
            step=operation,
            system_prompt=(
                "You are a novel-writing assistant. Produce the next revision of the "
                "attached manuscript as markdown. Return JSON with a single "
                "'chapter_markdown' string."
            ),
            user_prompt=(
                f"Operation: {operation}\n"
                f"Instruction: {instruction.strip()}\n\n"
                "Current manuscript:\n\n"
                f"{revision.content_markdown}"
            ),
            response_schema={"chapter_markdown": {"type": "string"}},
            metadata={
                "operation": operation,
                "document_id": revision.document_id,
                "base_revision_id": revision.id,
            },
        )
        generation_provider = self._ai_provider_factory(
            cast(TextGenerationProviderName, provider),
            model,
        )
        result = await generation_provider.generate_structured(task)
        proposal_markdown = result.content.get("chapter_markdown") or result.raw_text
        return _sanitize_chapter_markdown(proposal_markdown)

    async def generate_proposal(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
        *,
        operation: str,
        instruction: str,
        provider: str,
        model: str,
    ) -> tuple[str, str]:
        """Generate a proposal and return ``(proposal_markdown, base_revision_id)``."""
        _document, revision = self._load_revision(principal, project_id, document_id)
        proposal = await self._generate_proposal_text(
            revision,
            operation=operation,
            instruction=instruction,
            provider=provider,
            model=model,
        )
        return proposal, revision.id

    async def create_ai_proposal(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
        *,
        operation: str,
        instruction: str,
        provider: str = "mock",
        model: str = "studio-copilot-v1",
    ) -> dict[str, Any]:
        _document, revision = self._load_revision(principal, project_id, document_id)
        now = utcnow()
        try:
            proposal_markdown = _sanitize_chapter_markdown(
                await self._generate_proposal_text(
                    revision,
                    operation=operation,
                    instruction=instruction,
                    provider=provider,
                    model=model,
                )
            )
        except Exception as exc:
            return self._persist_failed_ai_job(
                project_id=project_id,
                document_id=document_id,
                operation=operation,
                provider=provider,
                model=model,
                instruction=instruction,
                base_revision_id=revision.id,
                error=str(exc),
                now=now,
            )
        return self._persist_completed_ai_job(
            project_id=project_id,
            document_id=document_id,
            operation=operation,
            provider=provider,
            model=model,
            instruction=instruction,
            base_revision_id=revision.id,
            proposal_markdown=proposal_markdown,
            now=now,
        )

    def _persist_failed_ai_job(
        self,
        *,
        project_id: str,
        document_id: str,
        operation: str,
        provider: str,
        model: str,
        instruction: str,
        base_revision_id: str,
        error: str,
        now: datetime,
    ) -> dict[str, Any]:
        job = self._repository.create_job(
            project_id=project_id,
            document_id=document_id,
            kind="proposal",
            operation=operation,
            status="failed",
            provider=provider,
            model=model,
            request_json=dump_json(
                {
                    "operation": operation,
                    "instruction": instruction,
                    "base_revision_id": base_revision_id,
                }
            ),
            result_json=dump_json(
                {
                    "proposal_markdown": "",
                    "base_revision_id": base_revision_id,
                    "accepted_revision_id": None,
                }
            ),
            error=error,
            retry_of_job_id=None,
            now=now,
        )
        self._repository.add_job_event(
            job.id,
            status="failed",
            details_json=dump_json({"error": error}),
            now=now,
        )
        return _job_payload(job)

    def _persist_completed_ai_job(
        self,
        *,
        project_id: str,
        document_id: str,
        operation: str,
        provider: str,
        model: str,
        instruction: str,
        base_revision_id: str,
        proposal_markdown: str,
        now: datetime,
    ) -> dict[str, Any]:
        job = self._repository.create_job(
            project_id=project_id,
            document_id=document_id,
            kind="proposal",
            operation=operation,
            status="completed",
            provider=provider,
            model=model,
            request_json=dump_json(
                {
                    "operation": operation,
                    "instruction": instruction,
                    "base_revision_id": base_revision_id,
                }
            ),
            result_json=dump_json(
                {
                    "proposal_markdown": proposal_markdown,
                    "base_revision_id": base_revision_id,
                    "accepted_revision_id": None,
                }
            ),
            error=None,
            retry_of_job_id=None,
            now=now,
        )
        self._repository.add_job_event(
            job.id,
            status="completed",
            details_json=dump_json({"proposal_only": True}),
            now=now,
        )
        self._repository.add_usage_event(
            project_id=project_id,
            job_id=job.id,
            provider=provider,
            model=model,
            prompt_tokens=_word_count(instruction),
            completion_tokens=_word_count(proposal_markdown),
            request_evidence_json=dump_json(
                {"operation": operation, "base_revision_id": base_revision_id}
            ),
            now=now,
        )
        return _job_payload(job)

    def accept_ai_proposal(
        self,
        principal: Principal,
        project_id: str,
        job_id: str,
    ) -> dict[str, Any]:
        owner_id, guest_session_id = _owner_scopes(principal)
        job = self._repository.get_job(
            project_id,
            job_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        if job.kind != "proposal" or job.document_id is None:
            raise NotFound("AI proposal not found.")
        result = cast(dict[str, Any], load_json(job.result_json))
        request = cast(dict[str, Any], load_json(job.request_json))
        if result.get("accepted_revision_id"):
            return _job_payload(job)
        document_id = job.document_id
        proposal = str(result.get("proposal_markdown", ""))
        base_revision_id = cast(str | None, request.get("base_revision_id"))

        document_service = DocumentService(self._repository)
        saved = document_service.save_document(
            principal,
            project_id,
            document_id,
            content_markdown=proposal,
            base_revision_id=base_revision_id,
            metadata={"ai_job_id": job_id},
            source="ai-accepted",
        )

        result["accepted_revision_id"] = saved["current_revision_id"]
        updated = self._repository.update_job(
            job_id,
            status=job.status,
            result_json=dump_json(result),
            now=utcnow(),
        )
        return _job_payload(updated)


class JobService:
    """Durable job execution, retry, and history."""

    def __init__(
        self,
        repository: StudioRepository,
        ai_service: AIService,
        review_service: ReviewService,
        export_service: ExportService,
    ) -> None:
        self._repository = repository
        self._ai_service = ai_service
        self._review_service = review_service
        self._export_service = export_service

    def list_jobs(self, principal: Principal, project_id: str) -> list[dict[str, Any]]:
        owner_id, guest_session_id = _owner_scopes(principal)
        jobs = self._repository.list_jobs(
            project_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return [_job_payload(job) for job in jobs]

    async def retry_job(
        self,
        principal: Principal,
        project_id: str,
        job_id: str,
    ) -> dict[str, Any]:
        owner_id, guest_session_id = _owner_scopes(principal)
        original = self._repository.get_job(
            project_id,
            job_id,
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        if original.status not in {"failed", "interrupted"}:
            raise InvalidOperation("Only failed or interrupted jobs may be retried.")
        now = utcnow()
        retry = self._repository.create_job(
            project_id=original.project_id,
            document_id=original.document_id,
            kind=original.kind,
            operation=original.operation,
            status="running",
            provider=original.provider,
            model=original.model,
            request_json=original.request_json,
            result_json="{}",
            error=None,
            retry_of_job_id=original.id,
            now=now,
        )
        self._repository.add_job_event(
            retry.id,
            status="running",
            details_json=dump_json({"retry_of": original.id}),
            now=now,
        )

        try:
            if retry.kind == "proposal":
                return await self._retry_ai_job(principal, retry)
            if retry.kind == "review":
                return await self._retry_review_job(principal, retry)
            if retry.kind == "export":
                return await self._retry_export_job(principal, retry)
            raise InvalidOperation(f"Unsupported job kind for retry: {retry.kind}")
        except Exception as exc:
            return self._fail_retry(retry.id, str(exc))

    async def _retry_ai_job(
        self,
        principal: Principal,
        retry: JobDto,
    ) -> dict[str, Any]:
        request = cast(dict[str, Any], load_json(retry.request_json))
        base_revision_id = cast(str | None, request.get("base_revision_id"))
        instruction = str(request.get("instruction", ""))
        if not base_revision_id or retry.document_id is None:
            raise InvalidOperation("Original AI job is missing base_revision_id.")
        proposal_markdown, generated_base_revision_id = await self._ai_service.generate_proposal(
            principal,
            retry.project_id,
            retry.document_id,
            operation=retry.operation,
            instruction=instruction,
            provider=retry.provider,
            model=retry.model,
        )
        now = utcnow()
        updated = self._repository.update_job(
            retry.id,
            status="completed",
            result_json=dump_json(
                {
                    "proposal_markdown": proposal_markdown,
                    "base_revision_id": generated_base_revision_id,
                    "accepted_revision_id": None,
                }
            ),
            finished_at=now,
            now=now,
        )
        self._repository.add_job_event(
            retry.id,
            status="completed",
            details_json=dump_json({"proposal_only": True}),
            now=now,
        )
        self._repository.add_usage_event(
            project_id=retry.project_id,
            job_id=retry.id,
            provider=retry.provider,
            model=retry.model,
            prompt_tokens=_word_count(instruction),
            completion_tokens=_word_count(proposal_markdown),
            request_evidence_json=dump_json(
                {"operation": retry.operation, "base_revision_id": generated_base_revision_id}
            ),
            now=now,
        )
        return _job_payload(updated)

    async def _retry_review_job(
        self,
        principal: Principal,
        retry: JobDto,
    ) -> dict[str, Any]:
        review = self._review_service.review_project(
            principal,
            retry.project_id,
            provider=retry.provider or "deterministic",
            model=retry.model or "studio-review-v1",
        )
        now = utcnow()
        updated = self._repository.update_job(
            retry.id,
            status="completed",
            result_json=dump_json(
                {
                    "review_id": review["id"],
                    "snapshot_id": review["snapshot_id"],
                    "summary": review["summary"],
                }
            ),
            finished_at=now,
            now=now,
        )
        self._repository.add_job_event(
            retry.id,
            status="completed",
            details_json=dump_json({"review_id": review["id"]}),
            now=now,
        )
        return _job_payload(updated)

    async def _retry_export_job(
        self,
        principal: Principal,
        retry: JobDto,
    ) -> dict[str, Any]:
        export = self._export_service.export_project(
            principal,
            retry.project_id,
            export_format=cast(ExportFormat, retry.operation),
        )
        now = utcnow()
        updated = self._repository.update_job(
            retry.id,
            status="completed",
            result_json=dump_json(
                {
                    "export_id": export["id"],
                    "snapshot_id": export["snapshot_id"],
                    "format": export["format"],
                    "download_url": export["download_url"],
                }
            ),
            finished_at=now,
            now=now,
        )
        self._repository.add_job_event(
            retry.id,
            status="completed",
            details_json=dump_json({"export_id": export["id"]}),
            now=now,
        )
        return _job_payload(updated)

    def _fail_retry(self, retry_id: str, error_message: str) -> dict[str, Any]:
        now = utcnow()
        updated = self._repository.update_job(
            retry_id,
            status="failed",
            error=error_message,
            finished_at=now,
            now=now,
        )
        self._repository.add_job_event(
            retry_id,
            status="failed",
            details_json=dump_json({"error": error_message}),
            now=now,
        )
        return _job_payload(updated)


class ImportService:
    """Legacy file-based workspace import."""

    def __init__(
        self,
        repository: StudioRepository,
        project_service: ProjectService,
        document_service: DocumentService,
    ) -> None:
        self._repository = repository
        self._project_service = project_service
        self._document_service = document_service

    def preview_legacy_workspace(self, source: Path) -> dict[str, Any]:
        source = source.expanduser().resolve()
        story_path = source / "story.yaml"
        if not story_path.is_file():
            raise InvalidOperation("Legacy workspace must contain story.yaml.")
        story = yaml.safe_load(story_path.read_text(encoding="utf-8")) or {}
        chapter_dir = source / "manuscript" / "chapters"
        chapters = sorted(chapter_dir.glob("chapter-*.md")) if chapter_dir.exists() else []
        source_hash = self._legacy_hash(source, [story_path, *chapters])
        return {
            "source": str(source),
            "source_hash": source_hash,
            "title": str(story.get("title", source.name)),
            "description": str(story.get("premise", "")),
            "chapter_count": len(chapters),
            "chapters": [
                {"filename": chapter.name, "bytes": chapter.stat().st_size}
                for chapter in chapters
            ],
        }

    def import_legacy_workspace(
        self,
        principal: Principal,
        source: Path,
    ) -> dict[str, Any]:
        preview = self.preview_legacy_workspace(source)
        existing = self._repository.find_project_by_import_hash(preview["source_hash"])
        if existing is not None:
            owner_id, guest_session_id = _owner_scopes(principal)
            project = self._repository.get_project(
                existing.id,
                owner_id=owner_id,
                guest_session_id=guest_session_id,
            )
            return _project_payload(project)
        new_project = self._project_service.create_project(
            principal,
            title=str(preview["title"]),
            description=str(preview["description"]),
            create_seed=False,
        )
        source_root = Path(str(preview["source"]))
        for position, chapter_path in enumerate(
            sorted((source_root / "manuscript" / "chapters").glob("chapter-*.md")),
            start=1,
        ):
            self._document_service.create_document(
                principal,
                new_project["id"],
                kind="chapter",
                title=f"Chapter {position}",
                content_markdown=chapter_path.read_text(encoding="utf-8"),
                position=position,
                metadata={"legacy_filename": chapter_path.name},
            )
        self._repository.set_project_import_hash(
            new_project["id"], str(preview["source_hash"])
        )
        owner_id, guest_session_id = _owner_scopes(principal)
        stored = self._repository.get_project(
            new_project["id"],
            owner_id=owner_id,
            guest_session_id=guest_session_id,
        )
        return _project_payload(stored)

    @staticmethod
    def _legacy_hash(root: Path, files: Iterable[Path]) -> str:
        digest = hashlib.sha256()
        digest.update(str(root).encode("utf-8"))
        for path in sorted(files):
            digest.update(path.relative_to(root).as_posix().encode("utf-8"))
            digest.update(path.read_bytes())
        return digest.hexdigest()


# ---------------------------------------------------------------------------
# Facade
# ---------------------------------------------------------------------------
class StudioStore:
    """Backward-compatible facade over the granular Studio domain services."""

    def __init__(
        self,
        repository: StudioRepository | None = None,
        *,
        data_dir: Path | None = None,
        ai_provider_factory: TextGenerationProviderFactory | None = None,
    ) -> None:
        self.repository = repository
        self.data_dir = data_dir
        self.ai_provider_factory = ai_provider_factory
        self._build_services()

    def _build_services(self) -> None:
        repository = self.repository
        data_dir = self.data_dir
        ai_factory = self.ai_provider_factory
        if repository is None:
            self.auth: AuthService | None = None
            self.project_service: ProjectService | None = None
            self.document_service: DocumentService | None = None
            self.revision_service: RevisionService | None = None
            self.snapshot_service: SnapshotService | None = None
            self.review_service: ReviewService | None = None
            self.export_service: ExportService | None = None
            self.ai_service: AIService | None = None
            self.job_service: JobService | None = None
            self.import_service: ImportService | None = None
        else:
            self.auth = AuthService(repository)
            self.project_service = ProjectService(repository)
            self.document_service = DocumentService(repository)
            self.revision_service = RevisionService(repository, self.document_service)
            self.snapshot_service = SnapshotService(repository)
            self.review_service = ReviewService(repository)
            if data_dir is None:
                raise InvalidOperation("data_dir is required when a repository is provided.")
            self.export_service = ExportService(repository, data_dir=data_dir)
            if ai_factory is None:
                raise InvalidOperation(
                    "ai_provider_factory is required when a repository is provided."
                )
            self.ai_service = AIService(repository, ai_factory)
            self.job_service = JobService(
                repository,
                self.ai_service,
                self.review_service,
                self.export_service,
            )
            self.import_service = ImportService(
                repository,
                self.project_service,
                self.document_service,
            )

    @property
    def database(self) -> Any:
        """Expose the underlying infrastructure database for diagnostics."""
        if self.repository is None:
            raise RuntimeError("StudioStore has not been configured with a repository.")
        return self.repository.database

    def _service(self, service: T | None) -> T:
        if service is None:
            raise RuntimeError("StudioStore has not been configured with a repository.")
        return service

    # Owner / auth
    def owner_exists(self) -> bool:
        return self._service(self.auth).owner_exists()

    def owner_principal(self, username: str | None = None) -> Principal:
        return self._service(self.auth).owner_principal(username)

    def setup_owner(self, username: str, password: str) -> dict[str, Any]:
        return self._service(self.auth).setup_owner(username, password)

    def create_owner_session(
        self,
        username: str,
        password: str,
    ) -> tuple[str, str, Principal]:
        return self._service(self.auth).create_owner_session(username, password)

    def create_guest_session(self) -> tuple[str, str, Principal]:
        return self._service(self.auth).create_guest_session()

    def csrf_token_for_session(self, token_hash: str) -> str | None:
        return self._service(self.auth).csrf_token_for_session(token_hash)

    def principal_from_token(self, token: str | None) -> Principal | None:
        return self._service(self.auth).principal_from_token(token)

    def logout(self, token: str | None) -> None:
        return self._service(self.auth).logout(token)

    def cleanup_expired_guests(self) -> int:
        return self._service(self.auth).cleanup_expired_guests()

    # Projects
    def create_project(
        self,
        principal: Principal,
        *,
        title: str,
        description: str = "",
        create_seed: bool = True,
    ) -> dict[str, Any]:
        return self._service(self.project_service).create_project(
            principal,
            title=title,
            description=description,
            create_seed=create_seed,
        )

    def list_projects(self, principal: Principal) -> list[dict[str, Any]]:
        return self._service(self.project_service).list_projects(principal)

    def get_project(self, principal: Principal, project_id: str) -> dict[str, Any]:
        return self._service(self.project_service).get_project(principal, project_id)

    def update_project(
        self,
        principal: Principal,
        project_id: str,
        *,
        title: str | None = None,
        description: str | None = None,
        settings: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        return self._service(self.project_service).update_project(
            principal,
            project_id,
            title=title,
            description=description,
            settings=settings,
        )

    # Documents
    def create_document(
        self,
        principal: Principal,
        project_id: str,
        *,
        kind: DocumentKind,
        title: str,
        content_markdown: str = "",
        position: int | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        return self._service(self.document_service).create_document(
            principal,
            project_id,
            kind=kind,
            title=title,
            content_markdown=content_markdown,
            position=position,
            metadata=metadata,
        )

    def get_document(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
    ) -> dict[str, Any]:
        return self._service(self.document_service).get_document(
            principal, project_id, document_id
        )

    def save_document(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
        *,
        content_markdown: str,
        base_revision_id: str | None,
        title: str | None = None,
        metadata: dict[str, Any] | None = None,
        source: str = "author",
    ) -> dict[str, Any]:
        return self._service(self.document_service).save_document(
            principal,
            project_id,
            document_id,
            content_markdown=content_markdown,
            base_revision_id=base_revision_id,
            title=title,
            metadata=metadata,
            source=source,
        )

    def reorder_documents(
        self,
        principal: Principal,
        project_id: str,
        document_ids: list[str],
    ) -> list[dict[str, Any]]:
        return self._service(self.document_service).reorder_documents(
            principal, project_id, document_ids
        )

    def search(
        self,
        principal: Principal,
        project_id: str,
        query: str,
    ) -> list[dict[str, Any]]:
        return self._service(self.document_service).search(
            principal, project_id, query
        )

    # Revisions
    def list_revisions(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
    ) -> list[dict[str, Any]]:
        return self._service(self.revision_service).list_revisions(
            principal, project_id, document_id
        )

    def restore_revision(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
        revision_id: str,
        *,
        base_revision_id: str | None,
    ) -> dict[str, Any]:
        return self._service(self.revision_service).restore_revision(
            principal,
            project_id,
            document_id,
            revision_id,
            base_revision_id=base_revision_id,
        )

    # Snapshots
    def create_snapshot(
        self,
        principal: Principal,
        project_id: str,
        *,
        reason: str,
    ) -> dict[str, Any]:
        return self._service(self.snapshot_service).create_snapshot(
            principal, project_id, reason=reason
        )

    def list_snapshots(
        self,
        principal: Principal,
        project_id: str,
    ) -> list[dict[str, Any]]:
        return self._service(self.snapshot_service).list_snapshots(
            principal, project_id
        )

    # Reviews
    def review_project(
        self,
        principal: Principal,
        project_id: str,
        *,
        provider: str = "deterministic",
        model: str = "studio-review-v1",
    ) -> dict[str, Any]:
        return self._service(self.review_service).review_project(
            principal, project_id, provider=provider, model=model
        )

    def list_reviews(
        self,
        principal: Principal,
        project_id: str,
    ) -> list[dict[str, Any]]:
        return self._service(self.review_service).list_reviews(principal, project_id)

    # Exports
    def export_project(
        self,
        principal: Principal,
        project_id: str,
        *,
        export_format: ExportFormat,
    ) -> dict[str, Any]:
        return self._service(self.export_service).export_project(
            principal, project_id, export_format=export_format
        )

    def list_exports(
        self,
        principal: Principal,
        project_id: str,
    ) -> list[dict[str, Any]]:
        return self._service(self.export_service).list_exports(principal, project_id)

    def export_path(
        self,
        principal: Principal,
        project_id: str,
        export_id: str,
    ) -> Path:
        return self._service(self.export_service).export_path(
            principal, project_id, export_id
        )

    # AI proposals
    async def create_ai_proposal(
        self,
        principal: Principal,
        project_id: str,
        document_id: str,
        *,
        operation: str,
        instruction: str,
        provider: str = "mock",
        model: str = "studio-copilot-v1",
    ) -> dict[str, Any]:
        return await self._service(self.ai_service).create_ai_proposal(
            principal,
            project_id,
            document_id,
            operation=operation,
            instruction=instruction,
            provider=provider,
            model=model,
        )

    def accept_ai_proposal(
        self,
        principal: Principal,
        project_id: str,
        job_id: str,
    ) -> dict[str, Any]:
        return self._service(self.ai_service).accept_ai_proposal(
            principal, project_id, job_id
        )

    # Jobs
    def list_jobs(self, principal: Principal, project_id: str) -> list[dict[str, Any]]:
        return self._service(self.job_service).list_jobs(principal, project_id)

    async def retry_job(
        self,
        principal: Principal,
        project_id: str,
        job_id: str,
    ) -> dict[str, Any]:
        return await self._service(self.job_service).retry_job(
            principal, project_id, job_id
        )

    # Import
    def preview_legacy_workspace(self, source: Path) -> dict[str, Any]:
        return self._service(self.import_service).preview_legacy_workspace(source)

    def import_legacy_workspace(
        self,
        principal: Principal,
        source: Path,
    ) -> dict[str, Any]:
        return self._service(self.import_service).import_legacy_workspace(
            principal, source
        )

    # Legacy payload helpers used by tests and diagnostics.
    def _project_payload(
        self,
        project: ProjectDto,
        *,
        include_documents: bool = True,
    ) -> dict[str, Any]:
        return _project_payload(project, include_documents=include_documents)

    def _document_payload(self, document: DocumentDto) -> dict[str, Any]:
        return _document_payload(document)

    def _revision_payload(self, revision: RevisionDto) -> dict[str, Any]:
        return _revision_payload(revision)

    def _snapshot_payload(self, snapshot: SnapshotDto) -> dict[str, Any]:
        return _snapshot_payload(snapshot)

    def _review_payload(self, review: ReviewDto) -> dict[str, Any]:
        return _review_payload(review)

    def _job_payload(self, job: JobDto) -> dict[str, Any]:
        return _job_payload(job)


# ---------------------------------------------------------------------------
# Singleton facade used by routers and CLI. Wired at application startup.
# ---------------------------------------------------------------------------
class _StudioStoreProxy:
    """Mutable proxy that lets routers and CLI import ``studio_store`` early.

    The real ``StudioStore`` instance is injected at application startup so
    that modules importing the singleton do not have to be imported after
    infrastructure wiring.
    """

    def __init__(self) -> None:
        self._instance: StudioStore | None = None

    def configure(
        self,
        instance: StudioStore,
    ) -> None:
        """Attach the configured facade instance."""
        self._instance = instance

    def __getattr__(self, name: str) -> Any:
        if self._instance is None:
            raise RuntimeError("StudioStore has not been configured with a repository.")
        return getattr(self._instance, name)

    def __setattr__(self, name: str, value: Any) -> None:
        if name == "_instance":
            super().__setattr__(name, value)
            return
        if self._instance is None:
            raise RuntimeError("StudioStore has not been configured with a repository.")
        setattr(self._instance, name, value)


_studio_store_proxy = _StudioStoreProxy()


def configure_studio_store(instance: StudioStore) -> None:
    """Attach a configured facade to the module-level singleton."""
    _studio_store_proxy.configure(instance)


def is_studio_store_configured() -> bool:
    """Return whether the module-level singleton has been wired."""
    return _studio_store_proxy._instance is not None


# ``studio_store`` is annotated as ``StudioStore`` so that importers get
# precise type checking; at runtime it is the proxy above.
studio_store: StudioStore = cast(StudioStore, _studio_store_proxy)
